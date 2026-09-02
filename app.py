import io
import re
import pandas as pd
import streamlit as st
import pdfplumber
from docx import Document
from rapidfuzz import fuzz

st.set_page_config(page_title="Attendance Compiler", layout="wide")

st.title("📊 Multi-Subject Attendance Compiler")
st.write("Automatically extracts Subject Name, Roll No/ID, and Name from Word, Excel, and PDF files!")

# --- HELPER FUNCTIONS ---

def extract_subject_from_text(text, filename=""):
    """Extracts Subject Name directly from text header or filename."""
    match = re.search(r'SUBJECT\s*[:-]\s*([^\n\r\|]+)', text, re.IGNORECASE)
    if match:
        subject_str = match.group(1).strip()
        subject_str = re.sub(r'\s+', ' ', subject_str)
        if subject_str:
            return subject_str.upper()

    course_match = re.search(r'(COURSE|PAPER|MODULE)\s*[:-]\s*([^\n\r\|]+)', text, re.IGNORECASE)
    if course_match:
        return course_match.group(2).strip().upper()

    clean_file = re.sub(r'[^A-Za-z0-9]', ' ', filename.rsplit('.', 1)[0]).strip()
    return clean_file.upper() if clean_file else "UNKNOWN SUBJECT"

def extract_from_docx(file):
    doc = Document(file)
    rows = []
    full_text = ""

    for p in doc.paragraphs:
        if p.text.strip():
            full_text += p.text + "\n"

    for table in doc.tables:
        for row in table.rows:
            cell_texts = [cell.text.strip() for cell in row.cells]
            rows.append(cell_texts)
            full_text += " ".join(cell_texts) + "\n"

    return rows, full_text

def extract_from_pdf(file):
    rows = []
    full_text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            full_text += text + "\n"
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    if row:
                        rows.append([str(cell).strip() if cell is not None else "" for cell in row])
    return rows, full_text

def parse_to_dataframe(file):
    ext = file.name.lower().rsplit('.', 1)[-1]
    raw_rows, full_text = [], ""

    if ext in ['xlsx', 'xls']:
        df_raw = pd.read_excel(file)
        full_text = " ".join([str(val) for val in df_raw.values.ravel() if pd.notna(val)])
        subject_name = extract_subject_from_text(full_text, file.name)
        cleaned_df = clean_and_structure_df(df_raw)
        return cleaned_df, subject_name

    elif ext == 'docx':
        raw_rows, full_text = extract_from_docx(file)
    elif ext == 'pdf':
        raw_rows, full_text = extract_from_pdf(file)

    if not raw_rows:
        return None, None

    subject_name = extract_subject_from_text(full_text, file.name)
    df_raw = pd.DataFrame(raw_rows)
    cleaned_df = clean_and_structure_df(df_raw)

    return cleaned_df, subject_name

def clean_and_structure_df(df):
    """Dynamically locates ID, Name, Total Classes, and Attendance Present columns."""
    df_clean = pd.DataFrame()
    id_col, name_col, total_col, att_col = None, None, None, None

    for col in df.columns:
        col_str = df[col].astype(str).str.upper()
        if not id_col and col_str.str.contains(r'\b(ID|ROLL|ENROLLMENT|S\.?NO|REG)\b', regex=True).any():
            id_col = col
        elif not name_col and col_str.str.contains(r'\b(NAME|STUDENT)\b', regex=True).any():
            name_col = col
        elif not total_col and col_str.str.contains(r'\b(TOTAL|HELD|CLASSES)\b', regex=True).any():
            total_col = col
        elif not att_col and col_str.str.contains(r'\b(ATTENDANCE|PRESENT|ATTENDED)\b', regex=True).any():
            att_col = col

    cols = list(df.columns)
    if not id_col and len(cols) >= 1: id_col = cols[0]
    if not name_col and len(cols) >= 2: name_col = cols[1]
    if not total_col and len(cols) >= 3: total_col = cols[2]
    if not att_col and len(cols) >= 4: att_col = cols[3]

    df_clean['ID'] = df[id_col].astype(str).str.strip() if id_col is not None else ""
    df_clean['NAME'] = df[name_col].astype(str).str.strip().str.upper() if name_col is not None else ""
    df_clean['TOTAL'] = pd.to_numeric(df[total_col], errors='coerce').fillna(0) if total_col is not None else 0
    df_clean['ATTENDANCE'] = pd.to_numeric(df[att_col], errors='coerce').fillna(0) if att_col is not None else 0

    # Filter non-student rows
    df_clean = df_clean[~df_clean['NAME'].str.contains('NAME|STUDENT|SUBJECT|TOTAL|ATTENDANCE|SEMESTER', na=False)]
    df_clean = df_clean[~df_clean['ID'].str.contains('ID|ROLL|S.NO', na=False)]
    df_clean = df_clean[df_clean['NAME'] != '']
    df_clean = df_clean[df_clean['NAME'] != 'NAN']

    return df_clean

def match_and_aggregate(processed_files):
    """Merges data by Roll No or Fuzzy Name Match. Inserts 'N/A' for missing subjects."""
    master_dict = {}
    all_subjects = [item['subject'] for item in processed_files]

    for item in processed_files:
        df = item['data']
        sub_name = item['subject']

        tot_col = f'TOTAL ({sub_name})'
        att_col = f'ATTENDANCE ({sub_name})'

        for _, row in df.iterrows():
            curr_id = str(row['ID']).strip() if row['ID'] and str(row['ID']).strip() != 'NAN' else ""
            curr_name = str(row['NAME']).strip().upper()
            curr_tot = float(row['TOTAL'])
            curr_att = float(row['ATTENDANCE'])

            matched_key = None

            # 1. Match by Roll No / ID
            if curr_id:
                for key, rec in master_dict.items():
                    if rec['ID'] == curr_id:
                        matched_key = key
                        break

            # 2. Match by Name (Token Set Ratio handles initial/mid/last variations)
            if not matched_key and curr_name:
                for key, rec in master_dict.items():
                    score = fuzz.token_set_ratio(curr_name, rec['NAME'])
                    if score >= 85:
                        matched_key = key
                        break

            # Add or update master dict
            if matched_key:
                master_dict[matched_key][tot_col] = master_dict[matched_key].get(tot_col, 0) + curr_tot
                master_dict[matched_key][att_col] = master_dict[matched_key].get(att_col, 0) + curr_att
                if not master_dict[matched_key]['ID'] and curr_id:
                    master_dict[matched_key]['ID'] = curr_id
            else:
                new_key = curr_id if curr_id else curr_name
                master_dict[new_key] = {
                    'ID': curr_id,
                    'NAME': curr_name,
                    tot_col: curr_tot,
                    att_col: curr_att
                }

    final_records = []
    
    # Format missing subject values as 'N/A' and compute overall percentage safely
    for key, rec in master_dict.items():
        row_dict = {'ID': rec['ID'], 'NAME': rec['NAME']}
        
        overall_total = 0
        overall_att = 0
        has_any_subject = False

        for sub in all_subjects:
            tot_c = f'TOTAL ({sub})'
            att_c = f'ATTENDANCE ({sub})'

            if tot_c in rec and att_c in rec:
                tot_val = rec[tot_c]
                att_val = rec[att_c]
                row_dict[tot_c] = int(tot_val)
                row_dict[att_c] = int(att_val)
                
                overall_total += tot_val
                overall_att += att_val
                has_any_subject = True
            else:
                row_dict[tot_c] = "N/A"
                row_dict[att_c] = "N/A"

        row_dict['OVERALL_TOTAL'] = int(overall_total) if has_any_subject else "N/A"
        row_dict['OVERALL_ATTENDANCE'] = int(overall_att) if has_any_subject else "N/A"
        row_dict['OVERALL_%'] = round((overall_att / overall_total * 100), 2) if (has_any_subject and overall_total > 0) else "N/A"

        final_records.append(row_dict)

    final_df = pd.DataFrame(final_records)
    return final_df

# --- STREAMLIT UI ---

uploaded_files = st.file_uploader(
    "Upload Word (.docx), Excel (.xlsx), or PDF (.pdf)",
    type=["docx", "pdf", "xlsx", "xls"],
    accept_multiple_files=True
)

if uploaded_files:
    processed_data = []

    st.write("### 📌 Automatically Detected Subjects")

    for idx, file in enumerate(uploaded_files):
        cleaned_df, subject_name = parse_to_dataframe(file)

        if cleaned_df is not None and not cleaned_df.empty:
            col1, col2 = st.columns([2, 2])
            with col1:
                st.write(f"📄 **File:** `{file.name}`")
            with col2:
                edited_subj = st.text_input(f"Detected Subject ({idx+1})", value=subject_name, key=f"subj_{idx}")

            processed_data.append({
                'filename': file.name,
                'subject': edited_subj,
                'data': cleaned_df
            })
        else:
            st.error(f"Could not extract table data from `{file.name}`")

    if processed_data and st.button("🚀 Compile Attendance", type="primary"):
        st.subheader("📋 Compiled Multi-Subject Attendance Report")
        result_df = match_and_aggregate(processed_data)
        st.dataframe(result_df, use_container_width=True)

        excel_buffer = io.BytesIO()
        with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
            result_df.to_excel(writer, index=False, sheet_name='Compiled Attendance')

        st.download_button(
            label="📥 Download Compiled Excel Report",
            data=excel_buffer.getvalue(),
            file_name="Compiled_Attendance_Report.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
